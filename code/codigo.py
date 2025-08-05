import google.generativeai as genai
from dotenv import load_dotenv
import os
from docx import Document
from docx.shared import Pt
import json

# Carga el archivo .env
load_dotenv()

# Obtiene la clave API del entorno
api_key = os.getenv("API_KEY")
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash-latest")

# Cargar datos desde el JSON
with open("estructura_presupuesto.json", encoding="utf-8") as f:
    estructura = json.load(f)

# Crear documento Word
doc = Document()
doc.add_heading("Especificaciones Técnicas", 0)

# Recorrer las partidas
en_categoría = ""
for item in estructura:
    if item[2] == "categoria":
        en_categoría = item[1]
    elif item[2] == "partida":
        codigo, descripcion, _, unidad = item[:4]

        # Prompt para generar especificación técnica
        prompt = f"""
Eres un ingeniero civil. Redacta una especificación técnica profesional para la siguiente partida:

Categoría: {en_categoría}
Código: {codigo}
Descripción: {descripcion}
Unidad: {unidad}

Incluye:
- Descripción general de la actividad.
- Materiales requeridos.
- Procedimientos constructivos.
- Equipos necesarios.
- Criterios de medición y pago.

Redacta de forma técnica y profesional.
"""
        response = model.generate_content(prompt)

        doc.add_heading(f"{codigo} - {descripcion}", level=1)
        paragraph = doc.add_paragraph(response.text)
        paragraph.style.font.size = Pt(11)

# Guardar el documento
nombre_archivo = "especificaciones_tecnicas.docx"
doc.save(nombre_archivo)
print(f"Documento generado: {nombre_archivo}")
